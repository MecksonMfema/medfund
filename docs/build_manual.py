#!/usr/bin/env python3
"""Convert docs/medfund-platform-manual.md to medfund-platform-manual.docx.

Handles a constrained markdown subset:
- YAML front matter (extracts title/subtitle/author/date for a title page)
- Headings #, ##, ###, ####
- Paragraphs
- Bullet lists (- )
- Numbered lists (1. )
- GitHub tables (| ... |)
- Code blocks (``` ... ```)
- Blockquotes (> )
- Inline: **bold**, *italic*, `code`
- Horizontal rules (---)
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = Path(__file__).resolve().parent
SRC = HERE / "medfund-platform-manual.md"
OUT = HERE / "medfund-platform-manual.docx"


INLINE_PATTERN = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def add_inline_runs(paragraph, text):
    """Append runs that honour **bold**, *italic*, `code`."""
    parts = INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def parse_table(table_lines, doc):
    """Render a GitHub-flavoured pipe table into a docx table."""
    rows = [line.strip().strip("|").split("|") for line in table_lines]
    rows = [[c.strip() for c in row] for row in rows]
    # Drop the alignment row (---|---)
    if len(rows) >= 2 and all(set(c.replace("-", "").replace(":", "").strip()) == set() for c in rows[1]):
        header = rows[0]
        body = rows[2:]
    else:
        header = rows[0]
        body = rows[1:]

    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        add_inline_runs(p, cell_text)
        for run in p.runs:
            run.bold = True
        set_cell_background(cell, "1F4E78")
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row in enumerate(body):
        for j, cell_text in enumerate(row[: len(header)]):
            cell = table.rows[1 + i].cells[j]
            cell.text = ""
            add_inline_runs(cell.paragraphs[0], cell_text)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for level, size in [("Heading 1", 22), ("Heading 2", 16), ("Heading 3", 13), ("Heading 4", 12)]:
        if level in styles:
            styles[level].font.size = Pt(size)
            styles[level].font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
            styles[level].font.bold = True


def add_title_page(doc, frontmatter):
    title = frontmatter.get("title", "MedFund Platform Manual")
    subtitle = frontmatter.get("subtitle", "")
    author = frontmatter.get("author", "")
    date = frontmatter.get("date", "")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n\n")
    run.font.size = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.italic = True
        run.font.size = Pt(14)

    for line in (author, date):
        if line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(12)

    doc.add_page_break()


def parse_frontmatter(lines):
    if not lines or lines[0].rstrip() != "---":
        return {}, lines
    end = None
    for i in range(1, len(lines)):
        if lines[i].rstrip() == "---":
            end = i
            break
    if end is None:
        return {}, lines
    fm = {}
    for line in lines[1:end]:
        if ":" in line:
            k, v = line.split(":", 1)
            v = v.strip().strip('"')
            fm[k.strip()] = v
    return fm, lines[end + 1 :]


def build_toc_field(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Right-click and choose 'Update Field' in Word to populate the table of contents."
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)


def main():
    text = SRC.read_text(encoding="utf-8")
    lines = text.splitlines()
    fm, body_lines = parse_frontmatter(lines)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    configure_styles(doc)
    add_title_page(doc, fm)

    p = doc.add_paragraph()
    run = p.add_run("Table of contents")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    toc_para = doc.add_paragraph()
    build_toc_field(toc_para)
    doc.add_page_break()

    i = 0
    n = len(body_lines)
    while i < n:
        line = body_lines[i]
        stripped = line.rstrip()

        # Code block
        if stripped.startswith("```"):
            j = i + 1
            buf = []
            while j < n and not body_lines[j].rstrip().startswith("```"):
                buf.append(body_lines[j])
                j += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(buf))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            i = j + 1
            continue

        # Horizontal rule
        if stripped == "---":
            add_horizontal_rule(doc)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            heading = m.group(2).strip()
            doc.add_heading(heading, level=level)
            i += 1
            continue

        # Table
        if stripped.startswith("|") and stripped.endswith("|"):
            tbl_lines = []
            while i < n and body_lines[i].strip().startswith("|"):
                tbl_lines.append(body_lines[i])
                i += 1
            parse_table(tbl_lines, doc)
            doc.add_paragraph()
            continue

        # Bullet list
        if re.match(r"^\s*-\s+", line):
            while i < n and re.match(r"^\s*-\s+", body_lines[i]):
                bullet_text = re.sub(r"^\s*-\s+", "", body_lines[i])
                p = doc.add_paragraph(style="List Bullet")
                add_inline_runs(p, bullet_text)
                i += 1
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", line):
            while i < n and re.match(r"^\s*\d+\.\s+", body_lines[i]):
                item_text = re.sub(r"^\s*\d+\.\s+", "", body_lines[i])
                p = doc.add_paragraph(style="List Number")
                add_inline_runs(p, item_text)
                i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            buf = []
            while i < n and body_lines[i].strip().startswith(">"):
                buf.append(re.sub(r"^\s*>\s?", "", body_lines[i]))
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            run_text = " ".join(b.strip() for b in buf if b.strip())
            add_inline_runs(p, run_text)
            for run in p.runs:
                run.italic = True
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Paragraph (possibly multi-line)
        buf = [line.strip()]
        i += 1
        while i < n:
            nxt = body_lines[i]
            if not nxt.strip():
                break
            if re.match(r"^(#{1,4})\s+", nxt.strip()):
                break
            if nxt.strip().startswith("|"):
                break
            if re.match(r"^\s*-\s+", nxt):
                break
            if re.match(r"^\s*\d+\.\s+", nxt):
                break
            if nxt.strip().startswith(">"):
                break
            if nxt.strip().startswith("```"):
                break
            if nxt.strip() == "---":
                break
            buf.append(nxt.strip())
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(buf))

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
